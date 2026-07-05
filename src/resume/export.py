from pathlib import Path
from src.config import RESUMES_DIR


def export_resume(markdown_text: str, output_path: Path, fmt: str = "pdf") -> Path:
    fmt = fmt.lower()
    if fmt == "pdf":
        return _export_pdf(markdown_text, output_path)
    elif fmt == "docx":
        return _export_docx(markdown_text, output_path)
    else:
        raise ValueError(f"Unsupported export format: {fmt}")


def _export_pdf(markdown_text: str, output_path: Path) -> Path:
    from fpdf import FPDF

    output_path = output_path.with_suffix(".pdf")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Resume", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    for line in markdown_text.split("\n"):
        if line.startswith("### "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, line[4:].strip(), new_x="LMARGIN", new_y="NEXT")
        elif line.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.cell(0, 9, line[3:].strip(), new_x="LMARGIN", new_y="NEXT")
        elif line.startswith("# "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, line[2:].strip(), new_x="LMARGIN", new_y="NEXT")
        elif line.strip() == "":
            pdf.ln(2)
        else:
            pdf.set_font("Helvetica", "", 10)
            clean = line.strip().lstrip("-*+").strip()
            pdf.multi_cell(0, 5, clean)

    pdf.output(str(output_path))
    return output_path


def _export_docx(markdown_text: str, output_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    output_path = output_path.with_suffix(".docx")
    doc = Document()

    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(12)
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(13)
        elif stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(14)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    doc.save(str(output_path))
    return output_path
